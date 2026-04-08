from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from pathlib import Path
import re

tpl = Path(r"C:\Users\Qiang\.openclaw\media\outbound\42abaf42-c0fb-4d74-be24-460b98b5494c.docx")
src = Path(r"C:\Users\Qiang\.openclaw\workspace-autosavants\articles\研究生创新与生涯规划_模块化无人机商业模式\versions\研究生创新与生涯规划_模块化无人机商业模式_v1_source.txt")
out = Path(r"C:\Users\Qiang\.openclaw\workspace-autosavants\articles\研究生创新与生涯规划_模块化无人机商业模式\versions\研究生创新与生涯规划_模块化无人机商业模式_v1_提交稿.docx")
alias = Path(r"C:\Users\Qiang\.openclaw\workspace-autosavants\paper_assignment_graduate_innovation_v1.docx")

doc = Document(tpl)

# 页面边距
sec = doc.sections[0]
sec.top_margin = Cm(2.5)
sec.bottom_margin = Cm(2.0)
sec.left_margin = Cm(2.8)
sec.right_margin = Cm(2.7)

# 读入正文源
raw = src.read_text(encoding='utf-8')
lines = [ln.rstrip() for ln in raw.splitlines()]

# 找到模板中题目占位段落位置
start_idx = None
for i,p in enumerate(doc.paragraphs):
    t = (p.text or '').strip()
    if '题目（黑体小三）自拟' in t:
        start_idx = i
        break
if start_idx is None:
    start_idx = len(doc.paragraphs)

# 删除从start_idx开始到末尾段落
for _ in range(len(doc.paragraphs)-start_idx):
    p = doc.paragraphs[start_idx]._element
    p.getparent().remove(p)

# 添加段落工具
body_style = None
if '论文正文' in [s.name for s in doc.styles]:
    body_style = '论文正文'

def set_font(run, name_cn='宋体', name_en='Times New Roman', size=12, bold=False):
    run.font.name = name_en
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name_cn)
    run.font.size = Pt(size)
    run.bold = bold

def add_para(text, kind='body'):
    p = doc.add_paragraph('')
    if body_style:
        p.style = body_style
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    r = p.add_run(text)

    if kind == 'title':
        p.alignment = 1
        p.paragraph_format.first_line_indent = Cm(0)
        set_font(r, name_cn='黑体', size=16, bold=False)  # 小三
    elif kind == 'h1':
        p.alignment = 0
        p.paragraph_format.first_line_indent = Cm(0)
        set_font(r, name_cn='黑体', size=14, bold=False)  # 四号
    elif kind == 'h2':
        p.alignment = 0
        p.paragraph_format.first_line_indent = Cm(0)
        set_font(r, name_cn='黑体', size=12, bold=False)  # 小四
    elif kind == 'meta':
        p.alignment = 0
        p.paragraph_format.first_line_indent = Cm(0)
        set_font(r, name_cn='宋体', size=12, bold=False)
    else:
        p.alignment = 0
        set_font(r, name_cn='宋体', size=12, bold=False)
    return p

h1_pat = re.compile(r'^[一二三四五六七八九十]+、')
h2_pat = re.compile(r'^(\d+\.\d+)')

for ln in lines:
    t = ln.strip()
    if not t:
        doc.add_paragraph('')
        continue

    if t.startswith('题目：'):
        add_para(t.replace('题目：','',1).strip(), 'title')
    elif t in ('摘要：','关键词：'):
        add_para(t, 'h1')
    elif t.startswith('关键词：'):
        add_para(t, 'meta')
    elif h1_pat.match(t):
        add_para(t, 'h1')
    elif h2_pat.match(t):
        add_para(t, 'h2')
    else:
        add_para(t, 'body')

# 保存
out.parent.mkdir(parents=True, exist_ok=True)
doc.save(out)

a = Document(out)
text = ''.join([p.text for p in a.paragraphs])
chars = len(''.join(text.split()))

import shutil
shutil.copy2(out, alias)

print(f'OUT={out}')
print(f'ALIAS={alias}')
print(f'CHARS_NO_SPACE={chars}')
