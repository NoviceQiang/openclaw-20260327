from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from pathlib import Path
import re, shutil

tpl = Path(r"C:\Users\Qiang\.openclaw\media\outbound\42abaf42-c0fb-4d74-be24-460b98b5494c.docx")
src = Path(r"C:\Users\Qiang\.openclaw\workspace-autosavants\articles\自然辩证法_人工智能时代嵌入式方向何去何从\versions\自然辩证法_人工智能时代嵌入式方向何去何从_v1_source.txt")
out_dir = Path(r"C:\Users\Qiang\.openclaw\workspace-autosavants\articles\自然辩证法_人工智能时代嵌入式方向何去何从\versions")
out = out_dir / "自然辩证法_人工智能时代嵌入式方向何去何从_v1_提交稿.docx"
alias = Path(r"C:\Users\Qiang\.openclaw\workspace-autosavants\paper_dialectics_of_nature_v1.docx")

doc = Document(tpl)

# 页面边距
sec = doc.sections[0]
sec.top_margin = Cm(2.5)
sec.bottom_margin = Cm(2.0)
sec.left_margin = Cm(2.8)
sec.right_margin = Cm(2.7)

# 找到题目占位段落位置
start_idx = None
for i, p in enumerate(doc.paragraphs):
    t = (p.text or "").strip()
    if "题目（黑体小三）自拟" in t:
        start_idx = i
        break
if start_idx is None:
    start_idx = len(doc.paragraphs)

# 删除从start_idx到末尾
for _ in range(len(doc.paragraphs) - start_idx):
    p_elem = doc.paragraphs[start_idx]._element
    p_elem.getparent().remove(p_elem)

def set_run_font(run, name_cn="宋体", name_en="Times New Roman", size=12, bold=False):
    run.font.name = name_en
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name_cn)
    run.font.size = Pt(size)
    run.bold = bold

def add_para(doc, text, kind="body"):
    p = doc.add_paragraph()
    p.alignment = 0
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)

    r = p.add_run(text)
    if kind == "title":
        p.alignment = 1
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(12)
        set_run_font(r, "黑体", "Times New Roman", 16, False)
    elif kind == "h1":
        p.alignment = 0
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        set_run_font(r, "黑体", "Times New Roman", 14, False)
    elif kind == "h2":
        p.alignment = 0
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        set_run_font(r, "黑体", "Times New Roman", 12, False)
    elif kind == "abstract":
        p.paragraph_format.first_line_indent = Cm(0)
        set_run_font(r, "宋体", "Times New Roman", 12, False)
    elif kind == "ref":
        p.paragraph_format.first_line_indent = Cm(0)
        set_run_font(r, "宋体", "Times New Roman", 10.5, False)
    else:
        set_run_font(r, "宋体", "Times New Roman", 12, False)
    return p

raw = src.read_text(encoding="utf-8")
lines = [ln.rstrip() for ln in raw.splitlines()]

h1_pat = re.compile(r"^[一二三四五六七八九十]+、")
h2_pat = re.compile(r"^(\d+\.\d+)\s")
in_refs = False

for ln in lines:
    t = ln.strip()
    if not t:
        continue

    if t.startswith("题目："):
        add_para(doc, t.replace("题目：", "", 1).strip(), "title")
    elif t.startswith("摘要："):
        add_para(doc, t, "abstract")
    elif t.startswith("关键词："):
        add_para(doc, t, "abstract")
    elif h1_pat.match(t):
        in_refs = "参考文献" in t
        add_para(doc, t, "h1")
    elif h2_pat.match(t):
        add_para(doc, t, "h2")
    elif in_refs:
        add_para(doc, t, "ref")
    else:
        add_para(doc, t, "body")

out.parent.mkdir(parents=True, exist_ok=True)
doc.save(out)
shutil.copy2(out, alias)

# 统计
doc2 = Document(out)
full_text = "".join([p.text for p in doc2.paragraphs])
chars = len("".join(full_text.split()))
print(f"OUT={out}")
print(f"ALIAS={alias}")
print(f"CHARS_NO_SPACE={chars}")
